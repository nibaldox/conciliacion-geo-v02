"""Tests for core.report_generator — plots, pie charts, image zip, Word report."""
import io
import zipfile

import matplotlib
matplotlib.use("Agg")
import numpy as np
import pandas as pd
import pytest
from docx import Document

from core.calculo_tronadura import procesar_pozos
from core.compliance_status import (
    STATUS_BANCO_ADICIONAL,
    STATUS_CUMPLE,
    STATUS_EXTRA,
    STATUS_FUERA,
    STATUS_NO_CUMPLE,
)
from core.param_extractor import BenchParams, ExtractionResult, compare_design_vs_asbuilt
from core.report_generator import (
    create_compliance_pie_charts,
    create_section_plot,
    generate_section_images_zip,
    generate_word_report,
)
from core.section_cutter import SectionLine


class TestCreateSectionPlot:
    def test_basic_section_plot(self):
        distances = np.linspace(0, 50, 20)
        elevations = 100 - distances * 0.5
        buf = create_section_plot(
            params_design=None,
            params_topo=None,
            distances_d=distances,
            elevations_d=elevations,
            distances_t=distances,
            elevations_t=elevations - 0.5,
        )
        assert isinstance(buf, io.BytesIO)
        assert buf.getbuffer().nbytes > 0
        buf.seek(0)
        assert buf.read(8).startswith(b"\x89PNG\r\n\x1a\n")

    def test_section_plot_with_bench_params(self):
        from core.param_extractor import BenchParams, ExtractionResult
        distances = np.array([0, 10, 20, 30, 40])
        elevations = np.array([100, 95, 90, 85, 80])
        bench = BenchParams(
            bench_number=1,
            crest_elevation=100.0,
            crest_distance=0.0,
            toe_elevation=85.0,
            toe_distance=40.0,
            bench_height=15.0,
            face_angle=70.0,
            berm_width=8.0,
        )
        params = ExtractionResult(
            section_name="BENCH", sector="TEST",
            benches=[bench], inter_ramp_angle=65.0, overall_angle=55.0,
        )
        buf = create_section_plot(
            params_design=params,
            params_topo=params,
            distances_d=distances,
            elevations_d=elevations,
            distances_t=distances,
            elevations_t=elevations,
        )
        assert isinstance(buf, io.BytesIO)
        assert buf.getbuffer().nbytes > 0

    def test_section_plot_empty(self):
        buf = create_section_plot(
            params_design=None,
            params_topo=None,
            distances_d=np.array([]),
            elevations_d=np.array([]),
            distances_t=np.array([]),
            elevations_t=np.array([]),
        )
        assert isinstance(buf, io.BytesIO)


class TestCompliancePieCharts:
    def test_pie_chart_returns_figures(self):
        comparisons = [
            {"height_status": "CUMPLE"},
            {"height_status": "CUMPLE"},
            {"height_status": "FUERA DE TOLERANCIA"},
            {"height_status": "NO CUMPLE"},
        ]
        result = create_compliance_pie_charts(comparisons)
        assert result is not None

    def test_empty_comparisons(self):
        result = create_compliance_pie_charts([])
        assert result is not None

    def test_all_cumple(self):
        comparisons = [{"height_status": "CUMPLE"}] * 5
        result = create_compliance_pie_charts(comparisons)
        assert result is not None

    def test_mixed_statuses_with_berm(self):
        comparisons = [
            {"height_status": "CUMPLE", "angle_status": "FUERA DE TOLERANCIA", "berm_status": "NO CUMPLE"},
            {"height_status": "FUERA DE TOLERANCIA", "angle_status": "CUMPLE", "berm_status": "CUMPLE"},
        ]
        result = create_compliance_pie_charts(comparisons)
        assert result is not None


class TestGenerateSectionImagesZip:
    def test_zip_creation_with_minimal_data(self):
        all_data = [
            {
                "section_name": "SEC_A",
                "params_design": None,
                "params_topo": None,
                "profile_d": (np.linspace(0, 30, 10), 100 - np.linspace(0, 30, 10)),
                "profile_t": (np.linspace(0, 30, 10), 99 - np.linspace(0, 30, 10)),
                "comparisons": [],
            }
        ]
        result = generate_section_images_zip(all_data)
        assert result is not None
        assert hasattr(result, "read") or isinstance(result, (bytes, io.BytesIO))

    def test_zip_creation_with_sections_list(self):
        from core.section_cutter import SectionLine
        sec = SectionLine(
            name="SEC_X",
            origin=np.array([0.0, 0.0]),
            azimuth=0.0,
            length=50.0,
        )
        all_data = [
            {
                "section_name": "SEC_X",
                "params_design": None,
                "params_topo": None,
                "profile_d": ([0.0, 25.0, 50.0], [100.0, 95.0, 90.0]),
                "profile_t": ([0.0, 25.0, 50.0], [99.0, 94.5, 89.5]),
                "comparisons": [],
            }
        ]
        result = generate_section_images_zip(all_data, sections=[sec])
        assert result is not None

    def test_zip_with_multiple_sections(self):
        all_data = []
        for i in range(3):
            all_data.append({
                "section_name": f"SEC_{i}",
                "params_design": None,
                "params_topo": None,
                "profile_d": ([0.0, 10.0], [100.0, 90.0]),
                "profile_t": ([0.0, 10.0], [99.0, 89.0]),
                "comparisons": [],
            })
        result = generate_section_images_zip(all_data)
        assert result is not None

    def test_zip_with_plot_options(self):
        all_data = [
            {
                "section_name": "SEC_Y",
                "params_design": None,
                "params_topo": None,
                "profile_d": ([0.0, 50.0], [100.0, 90.0]),
                "profile_t": ([0.0, 50.0], [99.0, 89.0]),
                "comparisons": [],
            }
        ]
        result = generate_section_images_zip(
            all_data, plot_options={"show_areas": False, "show_reconciled": False},
        )
        assert result is not None

    def test_zip_filters_sections_by_filtered_comps(self):
        all_data = [
            {
                "section_name": "SEC_A",
                "params_design": None,
                "params_topo": None,
                "profile_d": ([0.0, 25.0, 50.0], [100.0, 95.0, 90.0]),
                "profile_t": ([0.0, 25.0, 50.0], [99.0, 94.0, 89.0]),
                "comparisons": [],
            },
            {
                "section_name": "SEC_B",
                "params_design": None,
                "params_topo": None,
                "profile_d": ([0.0, 25.0], [100.0, 95.0]),
                "profile_t": ([0.0, 25.0], [99.0, 94.0]),
                "comparisons": [],
            },
        ]
        filtered_comps = [{"section": "SEC_A", "bench_num": 1}]
        result = generate_section_images_zip(all_data, filtered_comps=filtered_comps)
        with zipfile.ZipFile(result) as zf:
            names = set(zf.namelist())
        assert "SEC_A.png" in names
        assert "SEC_B.png" not in names


_REPORT_TOLERANCES = {
    "bench_height": {"neg": 1.0, "pos": 1.5},
    "face_angle": {"neg": 5.0, "pos": 5.0},
    "berm_width": {"min": 6.0},
}


def _bench(num, crest_z, toe_z, crest_x, toe_x, face=70.0, berm=9.0):
    return BenchParams(
        bench_number=num,
        crest_elevation=float(crest_z),
        crest_distance=float(crest_x),
        toe_elevation=float(toe_z),
        toe_distance=float(toe_x),
        bench_height=float(abs(crest_z - toe_z)),
        face_angle=float(face),
        berm_width=float(berm),
    )


def _extraction(section_name, benches, sector="Test"):
    return ExtractionResult(section_name=section_name, sector=sector, benches=benches)


def _matched_comparisons(section_name="S-01"):
    design = _extraction(
        section_name,
        [_bench(1, 100, 85, 10, 20), _bench(2, 85, 70, 30, 40)],
    )
    topo = _extraction(
        section_name,
        [
            _bench(1, 99, 83, 11, 21, face=68.0, berm=8.5),
            _bench(2, 84, 69, 31, 41, face=72.0, berm=10.0),
        ],
    )
    return compare_design_vs_asbuilt(design, topo, _REPORT_TOLERANCES)


def _all_data_item(section_name="S-01"):
    d = np.linspace(0, 50, 25)
    return {
        "section_name": section_name,
        "params_design": None,
        "params_topo": None,
        "profile_d": (d, 100.0 - d * 0.5),
        "profile_t": (d, 99.0 - d * 0.5),
    }


def _blast_holes_df():
    raw = pd.DataFrame(
        [
            {
                "Latitud_Geo": 10.0,
                "Longitud_Geo": 0.0,
                "Nombre_Banco": 4000.0,
                "Inclinacion_real": 0.0,
                "Azimuth_real": 0.0,
                "longitud_real": 10.0,
                "Kilos_Cargados_real": 200.0,
            }
        ]
    )
    return procesar_pozos(raw)[0]


class TestCreateSectionPlotBranches:
    def test_section_plot_show_areas(self):
        distances = np.linspace(0, 50, 25)
        elevations_d = 100 - distances * 0.5
        buf = create_section_plot(
            params_design=None,
            params_topo=None,
            distances_d=distances,
            elevations_d=elevations_d,
            distances_t=distances,
            elevations_t=elevations_d + 1.5,
            plot_options={"show_areas": True},
        )
        assert buf.getbuffer().nbytes > 0

    def test_section_plot_show_semaphore(self):
        distances = np.linspace(0, 50, 25)
        elevations_d = 100 - distances * 0.5
        buf = create_section_plot(
            params_design=None,
            params_topo=None,
            distances_d=distances,
            elevations_d=elevations_d,
            distances_t=distances,
            elevations_t=elevations_d - 3.0,
            plot_options={"show_semaphore": True},
        )
        assert buf.getbuffer().nbytes > 0

    def test_section_plot_z_limits(self):
        distances = np.linspace(0, 50, 25)
        elevations = 100 - distances * 0.5
        buf = create_section_plot(
            params_design=None,
            params_topo=None,
            distances_d=distances,
            elevations_d=elevations,
            distances_t=distances,
            elevations_t=elevations - 0.5,
            plot_options={"z_limits": (80.0, 105.0)},
        )
        assert buf.getbuffer().nbytes > 0

    def test_section_plot_z_span(self):
        distances = np.linspace(0, 50, 25)
        elevations = 100 - distances * 0.5
        buf = create_section_plot(
            params_design=None,
            params_topo=None,
            distances_d=distances,
            elevations_d=elevations,
            distances_t=distances,
            elevations_t=elevations - 0.5,
            plot_options={"z_span": 25.0},
        )
        assert buf.getbuffer().nbytes > 0

    def test_section_plot_show_pozos(self):
        distances = np.linspace(0, 50, 25)
        elevations = 100 - distances * 0.5
        section = SectionLine(
            name="S-01",
            origin=np.array([0.0, 0.0]),
            azimuth=90.0,
            length=200.0,
        )
        buf = create_section_plot(
            params_design=None,
            params_topo=None,
            distances_d=distances,
            elevations_d=elevations,
            distances_t=distances,
            elevations_t=elevations - 0.5,
            plot_options={"show_pozos": True, "blast_tolerance": 15.0},
            section=section,
            df_pozos=_blast_holes_df(),
        )
        assert buf.getbuffer().nbytes > 0


class TestGenerateWordReport:
    def test_word_report_creates_file(self, tmp_path):
        comps = _matched_comparisons("S-01")
        all_data = [_all_data_item("S-01")]
        out = tmp_path / "report.docx"
        generate_word_report(comps, all_data, str(out))

        assert out.exists()
        assert out.stat().st_size > 0

        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Informe de Conciliación Geotécnica" in text
        assert "Resumen Ejecutivo" in text
        assert "Tabla Resumen de Cumplimiento por Perfil" in text
        assert any("S-01" in c.text for t in doc.tables for r in t.rows for c in r.cells)

    def test_word_report_with_project_info(self, tmp_path):
        comps = _matched_comparisons("S-01")
        all_data = [_all_data_item("S-01")]
        out = tmp_path / "report.docx"
        generate_word_report(
            comps,
            all_data,
            str(out),
            project_info={"project": "Mina Norte", "author": "Geo Eng"},
        )

        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Mina Norte" in text
        assert "Geo Eng" in text

    def test_word_report_empty_comparisons(self, tmp_path):
        out = tmp_path / "report.docx"
        generate_word_report([], [], str(out))

        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "No se encontraron resultados para reportar." in text
        assert "No hay datos de comparación disponibles." in text

    def test_word_report_with_pozos_and_sections(self, tmp_path):
        comps = _matched_comparisons("S-01")
        all_data = [_all_data_item("S-01")]
        section = SectionLine(
            name="S-01",
            origin=np.array([0.0, 0.0]),
            azimuth=90.0,
            length=200.0,
        )
        out = tmp_path / "report.docx"
        generate_word_report(
            comps,
            all_data,
            str(out),
            df_pozos=_blast_holes_df(),
            sections=[section],
        )

        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Análisis de Perforación y Tronadura" in text
        assert "Cruce de Desviaciones vs Carga de Explosivo" in text
        assert any(
            "Pozos Cercanos" in c.text for t in doc.tables for r in t.rows for c in r.cells
        )

    def test_word_report_multiple_sections_page_break(self, tmp_path):
        comps = []
        all_data = []
        sections = []
        for i in range(8):
            name = f"S-{i:02d}"
            comps.extend(_matched_comparisons(name))
            all_data.append(_all_data_item(name))
            sections.append(
                SectionLine(
                    name=name,
                    origin=np.array([0.0, 0.0]),
                    azimuth=90.0,
                    length=200.0,
                )
            )
        out = tmp_path / "report.docx"
        generate_word_report(comps, all_data, str(out), sections=sections)

        assert out.exists()
        assert out.stat().st_size > 0
        doc = Document(str(out))
        assert any("Detalle por Sección" in p.text for p in doc.paragraphs)

    def test_word_report_na_values_for_missing_real(self, tmp_path):
        comps = [
            {
                "type": "MATCH",
                "section": "S-01",
                "bench_num": 1,
                "level": "85",
                "section_score": 40.0,
                "section_status": STATUS_NO_CUMPLE,
                "height_status": STATUS_CUMPLE,
                "angle_status": STATUS_FUERA,
                "berm_status": STATUS_NO_CUMPLE,
                "height_real": None,
                "angle_real": None,
                "berm_real": None,
                "height_design": 15.0,
                "angle_design": 70.0,
                "berm_design": 9.0,
            },
            {
                "type": "EXTRA",
                "section": "S-01",
                "bench_num": 999,
                "level": "70",
                "section_score": 40.0,
                "section_status": STATUS_NO_CUMPLE,
                "height_status": STATUS_EXTRA,
                "angle_status": "-",
                "berm_status": STATUS_BANCO_ADICIONAL,
                "height_real": 16.0,
                "angle_real": 71.0,
                "berm_real": 8.0,
                "height_design": None,
                "angle_design": None,
                "berm_design": None,
            },
        ]
        out = tmp_path / "report.docx"
        generate_word_report(comps, [], str(out))

        doc = Document(str(out))
        all_cells = [
            c.text for t in doc.tables for r in t.rows for c in r.cells
        ]
        assert "N/A" in all_cells
        assert "B999 (70)" in all_cells