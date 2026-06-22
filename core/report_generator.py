
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import numpy as np

from core.blast_correlation import (
    compute_blast_geotech_correlation,
    compute_pasadura_stats,
)
from core.calculo_tronadura import proyectar_pozos_en_seccion
from core.compliance_status import (
    STATUS_CUMPLE,
    STATUS_FUERA,
    STATUS_NO_CUMPLE,
)
from core.config import DEFAULTS

def create_section_plot(params_design, params_topo, distances_d, elevations_d, distances_t, elevations_t,
                        plot_options=None, section=None, df_pozos=None, filtered_bench_nums=None):
    if plot_options is None:
        plot_options = {}
    show_reconciled = plot_options.get('show_reconciled', True)
    show_areas = plot_options.get('show_areas', False)
    show_semaphore = plot_options.get('show_semaphore', False)
    show_pozos = plot_options.get('show_pozos', False)
    blast_tolerance = plot_options.get('blast_tolerance', 10.0)
    grid_height = plot_options.get('grid_height', 15.0)
    grid_ref = plot_options.get('grid_ref', 0.0)
    tolerances = plot_options.get('tolerances', {})
    if not tolerances:
        tolerances = {'bench_height': {'pos': 1.5}}

    fig, ax = plt.subplots(figsize=(10, 6))

    if len(distances_d) > 0:
        ax.plot(distances_d, elevations_d, color='royalblue', label='Diseño', linewidth=2)

    class ProfileHolder:
        def __init__(self, d, e):
            self.distances = np.array(d)
            self.elevations = np.array(e)

    pd_prof = ProfileHolder(distances_d, elevations_d)
    pt_prof = ProfileHolder(distances_t, elevations_t)

    if show_areas and len(distances_d) > 0 and len(distances_t) > 0:
        from core.geom_utils import calculate_area_between_profiles
        a_over, a_under, d_i, z_ref_i, z_eval_i = calculate_area_between_profiles(pd_prof, pt_prof)
        mask_u = z_eval_i >= z_ref_i
        if np.any(mask_u):
            ax.fill_between(d_i, z_ref_i, z_eval_i, where=mask_u, facecolor='blue', alpha=0.3)
        mask_o = z_eval_i < z_ref_i
        if np.any(mask_o):
            ax.fill_between(d_i, z_ref_i, z_eval_i, where=mask_o, facecolor='red', alpha=0.3)

    if len(distances_t) > 0:
        if show_semaphore:
            ax.plot(distances_t, elevations_t, color='gray', linewidth=0.5, alpha=0.7)
            from core.geom_utils import calculate_profile_deviation
            devs = calculate_profile_deviation(pd_prof, pt_prof)
            T = tolerances.get('bench_height', {}).get('pos', 1.5)

            mask_ok = devs <= T
            mask_warn = (devs > T) & (devs <= 1.5 * T)
            mask_nok = devs > 1.5 * T

            if np.any(mask_ok):
                ax.scatter(distances_t[mask_ok], elevations_t[mask_ok], color='#006100', s=10, zorder=5)
            if np.any(mask_warn):
                ax.scatter(distances_t[mask_warn], elevations_t[mask_warn], color='#FFD700', s=12, zorder=5)
            if np.any(mask_nok):
                ax.scatter(distances_t[mask_nok], elevations_t[mask_nok], color='#FF0000', s=12, zorder=5)

            ax.plot([], [], color='forestgreen', linewidth=2, label='Topografía Real')
        else:
            ax.plot(distances_t, elevations_t, color='forestgreen', label='Topografía Real', linewidth=2)

    from core.param_extractor import build_reconciled_profile_v2
    # Estilos por tipo de segmento para el perfil idealizado.
    # La cara del banco se dibuja continua; la berma (horizontal) con
    # guiones; la rampa (transición oblicua) con punteado fino.
    seg_style = {
        "face":      {"linestyle": "-",  "linewidth": 2.5},
        "berm_top":  {"linestyle": "--", "linewidth": 2.5},
        "ramp":      {"linestyle": ":",  "linewidth": 2.5},
    }
    if show_reconciled:
        if params_topo and params_topo.benches:
            rec_prof = build_reconciled_profile_v2(params_topo.benches, source="topo")
            if len(rec_prof.distances) > 0:
                # Trazo continuo base (para la leyenda y como fallback)
                ax.plot(rec_prof.distances, rec_prof.elevations,
                        color='#FF7F0E', label='Conciliado As-Built',
                        linewidth=2.5, zorder=4)
                # Sobre-grafico por segmento para diferenciar cara / berma / rampa
                for i in range(len(rec_prof.points) - 1):
                    p0 = rec_prof.points[i]
                    p1 = rec_prof.points[i + 1]
                    t = p0.segment_type
                    style = seg_style.get(t, seg_style["face"])
                    ax.plot([p0.distance, p1.distance],
                            [p0.elevation, p1.elevation],
                            color='#FF7F0E', **style, zorder=4)
                for b in params_topo.benches:
                    ax.plot(b.crest_distance, b.crest_elevation, marker='d', color='#FF7F0E', markersize=5, zorder=5)
                    ax.plot(b.toe_distance, b.toe_elevation, marker='d', color='#FF7F0E', markersize=5, zorder=5)

        if params_design and params_design.benches:
            rec_d_prof = build_reconciled_profile_v2(params_design.benches, source="design")
            if len(rec_d_prof.distances) > 0:
                ax.plot(rec_d_prof.distances, rec_d_prof.elevations,
                        color='royalblue', linestyle='--', linewidth=1.5, alpha=0.7)
                # Refuerzo de berma/rampa en diseño con línea punteada
                for i in range(len(rec_d_prof.points) - 1):
                    p0 = rec_d_prof.points[i]
                    p1 = rec_d_prof.points[i + 1]
                    t = p0.segment_type
                    if t == "ramp":
                        ax.plot([p0.distance, p1.distance],
                                [p0.elevation, p1.elevation],
                                color='royalblue', linestyle=':', linewidth=1.5, alpha=0.7)

    if params_topo and params_topo.benches:
        for bench in params_topo.benches:
            if filtered_bench_nums is not None and bench.bench_number not in filtered_bench_nums:
                continue
            ax.annotate(f"B{bench.bench_number}", xy=(bench.crest_distance, bench.crest_elevation),
                        xytext=(-8, 8), textcoords='offset points',
                        arrowprops=dict(arrowstyle="->", color="red", alpha=0.6),
                        fontsize=8, color="red")
            ax.annotate(f"Pa{bench.bench_number}", xy=(bench.toe_distance, bench.toe_elevation),
                        xytext=(8, -8), textcoords='offset points',
                        arrowprops=dict(arrowstyle="->", color="darkred", alpha=0.6),
                        fontsize=7, color="darkred")

    if show_pozos and df_pozos is not None and not df_pozos.empty and section is not None:
        projected = proyectar_pozos_en_seccion(
            df_pozos,
            origin=section.origin,
            azimuth=section.azimuth,
            length=section.length,
            tolerance=blast_tolerance,
        )
        if not projected.empty:
            for _, row in projected.iterrows():
                d_c = row['dist_along']
                d_t = row['dist_along_toe'] if 'dist_along_toe' in row else d_c
                z_c = row['Z_collar']
                z_t = row['Z_toe']
                ax.plot([d_c, d_t], [z_c, z_t], color='orange', linestyle='-', linewidth=1.5, alpha=0.5, zorder=3)
                ax.scatter(d_c, z_c, color='darkorange', s=15, zorder=4)

    all_d = np.concatenate([distances_d, distances_t])
    all_z = np.concatenate([elevations_d, elevations_t])
    valid_d = all_d[np.isfinite(all_d)]
    valid_z = all_z[np.isfinite(all_z)]

    if len(valid_d) > 0 and len(valid_z) > 0:
        xmin, xmax = float(np.min(valid_d)), float(np.max(valid_d))
        if 'z_span' in plot_options and plot_options['z_span'] is not None:
            z_span_val = plot_options['z_span']
            z_mid = (float(np.min(valid_z)) + float(np.max(valid_z))) / 2
            zmin = z_mid - z_span_val / 2
            zmax = z_mid + z_span_val / 2
        elif 'z_limits' in plot_options and plot_options['z_limits'] is not None:
            zmin, zmax = plot_options['z_limits']
        else:
            zmin, zmax = float(np.min(valid_z)), float(np.max(valid_z))
            z_pad = max((zmax - zmin) * 0.05, 5.0)
            zmin = zmin - z_pad
            zmax = zmax + z_pad

        min_x_span = (zmax - zmin) * 1.3
        current_x_span = xmax - xmin
        if current_x_span < min_x_span:
            x_mid = (xmin + xmax) / 2
            xmin = x_mid - min_x_span / 2
            xmax = x_mid + min_x_span / 2
        else:
            x_pad = max(current_x_span * 0.05, 5.0)
            xmin = xmin - x_pad
            xmax = xmax + x_pad

        ax.set_xlim(xmin, xmax)
        ax.set_ylim(zmin, zmax)

        if grid_height is not None and grid_height > 0:
            y_ticks = np.arange(np.floor((zmin - grid_ref) / grid_height) * grid_height + grid_ref,
                                np.ceil((zmax - grid_ref) / grid_height) * grid_height + grid_ref + grid_height,
                                grid_height)
            ax.set_yticks(y_ticks)

    title_name = "N/A"
    title_sector = "N/A"
    if section is not None:
        title_name = section.name
        title_sector = section.sector
    elif params_design is not None:
        title_name = params_design.section_name
        title_sector = params_design.sector
    elif params_topo is not None:
        title_name = params_topo.section_name
        title_sector = params_topo.sector

    ax.set_title(f"Sección: {title_name} - {title_sector}")
    ax.set_xlabel("Distancia (m)")
    ax.set_ylabel("Elevación (m)")
    ax.grid(True, linestyle='--', color='lightgray', alpha=0.7)
    ax.legend(loc='upper right', fontsize=7, labelspacing=0.2, handletextpad=0.3, borderaxespad=0.3)
    ax.set_aspect('equal', adjustable='box')

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def create_compliance_pie_charts(comparisons):
    """Create a row of donut charts (one per parameter) of compliance breakdown.

    Categories follow the tripartite status emitted by
    `param_extractor._evaluate_status` and `compare_design_vs_asbuilt`:
    CUMPLE / FUERA DE TOLERANCIA / NO CUMPLE. The `berm_status` is binary in
    the pipeline, so its FUERA DE TOLERANCIA slice is always zero.
    """
    keys = ['height_status', 'angle_status', 'berm_status']
    labels = ['Altura de Banco', 'Ángulo de Cara', 'Ancho de Berma']

    # Soft, high-contrast palette (readable on screen and printed).
    category_colors = {
        STATUS_CUMPLE:        '#7FBF7F',  # soft green
        STATUS_FUERA:         '#FFD27F',  # soft amber
        STATUS_NO_CUMPLE:     '#F08C8C',  # soft red
    }
    category_order = [STATUS_CUMPLE, STATUS_FUERA, STATUS_NO_CUMPLE]

    fig, axes = plt.subplots(1, 3, figsize=(12, 4))

    match_comps = [c for c in comparisons if c.get('type') == 'MATCH']
    for idx, (key, title) in enumerate(zip(keys, labels)):
        ax = axes[idx]
        counts = {cat: sum(1 for c in match_comps if c[key] == cat) for cat in category_order}

        labels_to_show, sizes, colors_to_show = [], [], []
        for cat in category_order:
            val = counts[cat]
            if val > 0:
                labels_to_show.append(f"{cat}\n({val})")
                sizes.append(val)
                colors_to_show.append(category_colors[cat])

        if sum(sizes) > 0:
            # Donut: pie with a hole in the middle shows the total and the
            # global % de logro, complementing the per-slice percentages.
            wedges, _texts, autotexts = ax.pie(
                sizes,
                labels=labels_to_show,
                autopct='%1.1f%%',
                startangle=90,
                colors=colors_to_show,
                wedgeprops={'width': 0.38, 'edgecolor': 'white', 'linewidth': 1.5},
                textprops={'fontsize': 9, 'weight': 'bold'},
            )
            for at in autotexts:
                at.set_color('black')
                at.set_fontsize(8)

            total = sum(sizes)
            n_ok = counts[STATUS_CUMPLE]
            pct = (n_ok / total * 100) if total > 0 else 0
            ax.text(0, 0, f"{pct:.0f}%\nLogro",
                    ha='center', va='center', fontsize=10, weight='bold')
            ax.set_title(title, fontsize=11, weight='bold', pad=10)
        else:
            ax.text(0.5, 0.5, 'Sin Datos', ha='center', va='center')
            ax.axis('off')

    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _setup_landscape_doc() -> Document:
    """Create a Document with landscape orientation and 0.5in margins."""
    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    return doc


def generate_word_report(comparisons, all_data, output_path, project_info=None,
                         df_pozos=None, sections=None, plot_options=None):
    if project_info is None:
        project_info = {}

    doc = _setup_landscape_doc()

    if plot_options is None:
        plot_options = {}

    title = doc.add_heading(f"Informe de Conciliación Geotécnica", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"Proyecto: {project_info.get('project', 'N/A')}\n").bold = True
    p.add_run(f"Elaborado por: {project_info.get('author', 'N/A')}\n")
    p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\n")

    doc.add_heading("1. Resumen Ejecutivo", level=1)

    if comparisons:
        match_comps = [c for c in comparisons if c.get('type') == 'MATCH']
        if match_comps:
            section_score = match_comps[0].get('section_score', 0.0)
            section_status = match_comps[0].get('section_status', STATUS_NO_CUMPLE)
        else:
            section_score = 0.0
            section_status = STATUS_NO_CUMPLE

        doc.add_paragraph(f"Se evaluaron {len(match_comps)} bancos emparejados.")
        doc.add_paragraph(f"Cumplimiento General (Ponderado): {section_score:.0f}/100 — {section_status}")

        try:
            pie_stream = create_compliance_pie_charts(comparisons)
            p_pie = doc.add_paragraph()
            p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_pie.add_run().add_picture(pie_stream, width=Inches(8.5))
            pie_stream.close()
        except Exception as e:
            doc.add_paragraph(f"(Error al generar gráficos de torta: {e})")

        doc.add_paragraph("Resumen por parámetro:")
        # Ternary table: CUMPLE + FUERA DE TOLERANCIA + NO CUMPLE = Total.
        # % Logro is CUMPLE / Total * 100. The last column shows the
        # average real value for the evaluated section (real height in m,
        # real face angle in degrees, real berm width in m).
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parámetro'
        hdr_cells[1].text = STATUS_CUMPLE
        hdr_cells[2].text = 'Fuera de Tolerancia'
        hdr_cells[3].text = STATUS_NO_CUMPLE
        hdr_cells[4].text = '% Logro'
        hdr_cells[5].text = 'Valor Promedio (Real)'

        for col_idx, h in enumerate(['Parámetro', STATUS_CUMPLE,
                                     'Fuera de Tolerancia', STATUS_NO_CUMPLE,
                                     '% Logro', 'Valor Promedio (Real)']):
            for paragraph in hdr_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9.5)

        # Each row aggregates over MATCH comparisons and shows the real
        # value average. Berm/angle/height use the *_real field of the
        # comparison dict; averages ignore None.
        param_avg_specs = [
            ('height_status', 'Altura',          'height_real', 'm'),
            ('angle_status',  'Ángulo Cara',     'angle_real',  '°'),
            ('berm_status',   'Berma',           'berm_real',   'm'),
        ]
        for key, label, real_field, unit in param_avg_specs:
            row_cells = table.add_row().cells
            n_ok = sum(1 for c in match_comps if c[key] == STATUS_CUMPLE)
            n_ft = sum(1 for c in match_comps if c[key] == STATUS_FUERA)
            n_nok = sum(1 for c in match_comps if c[key] == STATUS_NO_CUMPLE)
            total = n_ok + n_ft + n_nok
            pct = (n_ok / total * 100) if total > 0 else 0

            real_values = [c[real_field] for c in match_comps
                           if c.get(real_field) is not None]
            if real_values:
                avg_val = sum(real_values) / len(real_values)
                avg_str = f"{avg_val:.2f} {unit}"
            else:
                avg_str = "N/A"

            row_cells[0].text = label
            row_cells[1].text = str(n_ok)
            row_cells[2].text = str(n_ft)
            row_cells[3].text = str(n_nok)
            row_cells[4].text = f"{pct:.1f}%"
            row_cells[5].text = avg_str
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9.0)
    else:
        doc.add_paragraph("No se encontraron resultados para reportar.")

    doc.add_heading("2. Tabla Resumen de Cumplimiento por Perfil", level=1)
    if comparisons:
        # Per-profile compliance table. Previously the design/real values
        # were packed into a single "Dise / Real" cell. We now split them
        # into separate columns to make the report easier to scan. The
        # column order matches the previous layout (Diseño first, Real
        # second) for each parameter, so the table reads top-to-bottom as:
        # Sección · Banco · H. Dise · H. Real · Ang. Dise · Ang. Real ·
        # Berma Dise · Berma Real.
        table_summary = doc.add_table(rows=1, cols=8)
        table_summary.style = 'Table Grid'

        headers = [
            'Sección', 'Banco (cota)',
            'H. Diseño (m)', 'H. Real (m)',
            'Ang. Diseño (°)', 'Ang. Real (°)',
            'Berma Diseño (m)', 'Berma Real (m)',
        ]
        for col_idx, h in enumerate(headers):
            hdr_cell = table_summary.rows[0].cells[col_idx]
            hdr_cell.text = h
            for paragraph in hdr_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9.0)

        for c in comparisons:
            row_cells = table_summary.add_row().cells
            row_cells[0].text = c.get('section', 'N/A')

            b_num = c.get('bench_num', 'N/A')
            b_level = c.get('level', 'N/A')
            row_cells[1].text = f"B{b_num} ({b_level})"

            row_cells[2].text = (f"{c['height_design']:.1f}"
                                 if c.get('height_design') is not None else "N/A")
            row_cells[3].text = (f"{c['height_real']:.1f}"
                                 if c.get('height_real') is not None else "N/A")

            row_cells[4].text = (f"{c['angle_design']:.1f}"
                                 if c.get('angle_design') is not None else "N/A")
            row_cells[5].text = (f"{c['angle_real']:.1f}"
                                 if c.get('angle_real') is not None else "N/A")

            row_cells[6].text = (f"{c['berm_design']:.1f}"
                                 if c.get('berm_design') is not None else "N/A")
            row_cells[7].text = (f"{c['berm_real']:.1f}"
                                 if c.get('berm_real') is not None else "N/A")

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
    else:
        doc.add_paragraph("No hay datos de comparación disponibles.")

    doc.add_page_break()

    doc.add_heading("3. Detalle por Sección", level=1)

    valid_items = []
    for item in all_data:
        sec_name = item['section_name']
        sec_comps = [c for c in comparisons if c['section'] == sec_name]
        if sec_comps:
            valid_items.append((item, sec_comps))

    chunks = [valid_items[i:i + 6] for i in range(0, len(valid_items), 6)]

    for chunk_idx, chunk in enumerate(chunks):
        n_rows = (len(chunk) + 2) // 3
        table_plots = doc.add_table(rows=n_rows, cols=3)
        table_plots.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_plots.autofit = False

        for col_idx in range(3):
            table_plots.columns[col_idx].width = Inches(3.2)

        for idx, (item, sec_comps) in enumerate(chunk):
            row_idx = idx // 3
            col_idx = idx % 3
            sec_name = item['section_name']
            pd = item['params_design']
            pt = item['params_topo']
            prof_d = item['profile_d']
            prof_t = item['profile_t']

            sec_obj = None
            if sections:
                for s in sections:
                    if s.name == sec_name:
                        sec_obj = s
                        break

            filtered_bench_nums = {c['bench_num'] for c in sec_comps}

            img_stream = create_section_plot(
                pd, pt, prof_d[0], prof_d[1], prof_t[0], prof_t[1],
                plot_options=plot_options, section=sec_obj, df_pozos=df_pozos,
                filtered_bench_nums=filtered_bench_nums
            )

            cell = table_plots.rows[row_idx].cells[col_idx]
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cell.paragraph_format.space_before = Pt(0)
            p_cell.paragraph_format.space_after = Pt(0)
            p_cell.add_run().add_picture(img_stream, width=Inches(3.0))
            img_stream.close()

        if chunk_idx < len(chunks) - 1:
            doc.add_page_break()

    if df_pozos is not None and not df_pozos.empty:
        doc.add_page_break()
        doc.add_heading("4. Análisis de Perforación y Tronadura", level=1)

        stats = compute_pasadura_stats(df_pozos)
        p_mean = stats["mean"]
        p_optimal = stats["optimal_count"]
        p_pct = stats["optimal_pct"]

        p1 = doc.add_paragraph()
        p1.add_run("Estadísticas Generales de Perforación y Voladura:\n").bold = True
        p1.add_run(f"- Total de Pozos Registrados: {stats['total']}\n")
        p1.add_run(f"- Pasadura Promedio (Sub-drilling): {p_mean:.2f} m\n")
        p1.add_run(f"- Porcentaje de Pozos en Pasadura Óptima (0.5m a 1.5m): {p_pct:.1f}% ({p_optimal} pozos)\n")

        if sections and comparisons:
            doc.add_heading("Cruce de Desviaciones vs Carga de Explosivo", level=2)
            r_label = int(DEFAULTS.blast_correlation_radius_m)
            doc.add_paragraph(
                f"A continuación se detalla la cantidad de pozos y la carga de explosivo acumulada en un radio de {r_label} metros "
                "respecto al eje de cada sección transversal, cruzada con su respectiva desviación absoluta media:"
            )

            corr_rows = compute_blast_geotech_correlation(
                df_pozos=df_pozos,
                sections=sections,
                comparisons=comparisons,
            )
            if corr_rows:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'

                headers = ["Sección", "Pozos Cercanos", "Kilos de Explosivo", "Desviación Media (m)"]
                for idx, h in enumerate(headers):
                    table.rows[0].cells[idx].text = h

                for row in corr_rows:
                    sec_name, num_wells, total_kg, avg_dev = row.as_tuple()
                    row_cells = table.add_row().cells
                    row_cells[0].text = sec_name
                    row_cells[1].text = str(num_wells)
                    row_cells[2].text = f"{total_kg:.0f}"
                    row_cells[3].text = f"{avg_dev:.2f}"

    doc.save(output_path)


def generate_section_images_zip(all_data, plot_options=None, sections=None, df_pozos=None, filtered_comps=None):
    import zipfile

    if plot_options is None:
        plot_options = {}

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for item in all_data:
            sec_name = item['section_name']

            filtered_bench_nums = None
            if filtered_comps:
                sec_comps = [c for c in filtered_comps if c['section'] == sec_name]
                if not sec_comps:
                    continue
                filtered_bench_nums = {c['bench_num'] for c in sec_comps}

            pd = item['params_design']
            pt = item['params_topo']
            prof_d = item['profile_d']
            prof_t = item['profile_t']

            sec_obj = None
            if sections:
                for s in sections:
                    if s.name == sec_name:
                        sec_obj = s
                        break

            img_buf = create_section_plot(
                pd, pt, prof_d[0], prof_d[1], prof_t[0], prof_t[1],
                plot_options=plot_options, section=sec_obj, df_pozos=df_pozos,
                filtered_bench_nums=filtered_bench_nums
            )

            filename = f"{sec_name}.png"
            zip_file.writestr(filename, img_buf.getvalue())
            img_buf.close()

    zip_buffer.seek(0)
    return zip_buffer
